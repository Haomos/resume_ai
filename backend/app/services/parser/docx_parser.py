"""DOCX text extraction using python-docx."""

from io import BytesIO
from docx import Document


def extract_text(content: bytes) -> str | None:
    """Extract text from DOCX bytes. Returns None on failure or empty result."""
    try:
        doc = Document(BytesIO(content))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        return "\n\n".join(paragraphs) if paragraphs else None
    except Exception:
        return None


def extract_structured_html(content: bytes) -> str | None:
    """Extract structured HTML from DOCX, preserving headings / lists / paragraphs."""
    try:
        doc = Document(BytesIO(content))
        parts: list[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style = para.style.name.lower() if para.style and para.style.name else ""
            if "heading" in style or "标题" in style:
                level = "2"
                # Try to extract level number
                import re
                m = re.search(r"(\d+)", style)
                if m:
                    level = min(int(m.group(1)), 3)
                parts.append(f"<h{level}>{_escape_html(text)}</h{level}>")
            elif "list" in style or any(text.startswith(b) for b in ["•", "·", "-", "*", "・"]):
                parts.append(f"<li>{_escape_html(text.lstrip('•·-*・ '))}</li>")
            else:
                parts.append(f"<p>{_escape_html(text)}</p>")
        # Wrap consecutive <li> into <ul>
        result: list[str] = []
        in_list = False
        for part in parts:
            if part.startswith("<li>"):
                if not in_list:
                    result.append("<ul>")
                    in_list = True
                result.append(part)
            else:
                if in_list:
                    result.append("</ul>")
                    in_list = False
                result.append(part)
        if in_list:
            result.append("</ul>")
        return "\n".join(result) if result else None
    except Exception:
        text = extract_text(content)
        return None if text is None else f"<p>{_escape_html(text).replace(chr(10), '</p><p>')}</p>"


def _escape_html(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
