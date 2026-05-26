"""Tests for parser factory dispatch logic."""

import pytest
from io import BytesIO
from docx import Document

from app.services.parser.factory import extract_text
from app.services.parser.html_parser import extract_text as html_extract
from app.services.parser.pdf_parser import extract_text as pdf_extract
from app.services.parser.docx_parser import extract_text as docx_extract


def test_html_parser_basic():
    html = b"<html><body><p>Hello World</p></body></html>"
    text = html_extract(html)
    assert "Hello World" in text


def test_html_parser_script_stripped():
    html = b"""<html><head><script>alert(1);</script></head>
<body><p>Content</p></body></html>"""
    text = html_extract(html)
    assert "alert" not in text
    assert "Content" in text


def test_docx_parser_basic():
    doc = Document()
    doc.add_paragraph("Line one")
    doc.add_paragraph("Line two")
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    text = docx_extract(buf.read())
    assert "Line one" in text
    assert "Line two" in text


def test_pdf_parser_empty():
    # Empty bytes should return None gracefully
    assert pdf_extract(b"") is None


def test_factory_txt():
    assert extract_text(b"Hello txt", "resume.txt") == "Hello txt"


def test_factory_image_disabled_returns_none(monkeypatch):
    """Image extension is now routed to image_parser; with OCR disabled (default)
    it returns None — but the dispatch path itself is exercised."""
    monkeypatch.delenv("OCR_ENABLED", raising=False)
    from app.config import get_config
    get_config.cache_clear()
    # Use real PNG bytes so the decode step succeeds; disabled gate returns None
    from io import BytesIO
    from PIL import Image
    buf = BytesIO()
    Image.new("RGB", (10, 10), (255, 255, 255)).save(buf, format="PNG")
    assert extract_text(buf.getvalue(), "resume.png") is None


def test_factory_unknown_extension():
    """Truly unsupported extension → None."""
    assert extract_text(b"data", "resume.xyz") is None
